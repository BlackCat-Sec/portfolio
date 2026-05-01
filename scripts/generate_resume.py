from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ACCENT = RGBColor(15, 76, 129)
TEXT = RGBColor(27, 31, 36)
MUTED = RGBColor(87, 96, 106)


def set_page_margins(section):
    section.top_margin = Inches(0.48)
    section.bottom_margin = Inches(0.48)
    section.left_margin = Inches(0.58)
    section.right_margin = Inches(0.58)


def style_run(run, *, size=10.5, bold=False, color=TEXT, font_name="Arial"):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


def add_text(paragraph, text, **kwargs):
    run = paragraph.add_run(text)
    style_run(run, **kwargs)
    return run


def set_spacing(paragraph, *, before=0, after=0, line=1.0):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_section_heading(document, title):
    p = document.add_paragraph()
    set_spacing(p, before=7, after=4, line=1.0)
    add_text(p, title.upper(), size=10.5, bold=True, color=ACCENT)
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "B7C4D3")
    pBdr.append(bottom)
    return p


def add_body_paragraph(document, text, *, size=10.1, after=1.5, line=1.06, color=TEXT):
    p = document.add_paragraph()
    set_spacing(p, after=after, line=line)
    add_text(p, text, size=size, color=color)
    return p


def add_bullet(document, text, *, indent=0.18, size=9.9, after=0.5, line=1.03):
    p = document.add_paragraph()
    fmt = p.paragraph_format
    fmt.left_indent = Inches(indent)
    fmt.first_line_indent = Inches(-0.14)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    add_text(p, "- ", size=size, bold=True, color=ACCENT)
    add_text(p, text, size=size)
    return p


def add_project(document, title, stack, bullets):
    p = document.add_paragraph()
    set_spacing(p, after=1.3, line=1.04)
    add_text(p, title, size=10.6, bold=True)
    add_text(p, " | ", size=10.2, color=MUTED)
    add_text(p, stack, size=10.2, color=MUTED)
    for bullet in bullets:
        add_bullet(document, bullet)


def build_resume(output_path: Path):
    document = Document()
    section = document.sections[0]
    set_page_margins(section)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.1)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, after=1, line=1.0)
    add_text(p, "Mokshith H S", size=19.5, bold=True)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, after=1.5, line=1.0)
    add_text(
        p,
        "Cybersecurity Internship Candidate | SOC / Blue Team | Secure Software",
        size=10.5,
        bold=True,
        color=ACCENT,
    )

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, after=0.5, line=1.0)
    add_text(
        p,
        "gowdas573201@gmail.com | 6360571927 | Hassan, Karnataka, India",
        size=9.5,
        color=MUTED,
    )

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, after=4, line=1.0)
    add_text(
        p,
        "linkedin.com/in/mokshith-hs | github.com/BlackCat-Sec | blackcat-sec.github.io/portfolio",
        size=9.5,
        color=MUTED,
    )

    add_section_heading(document, "Professional Summary")
    add_body_paragraph(
        document,
        "Cybersecurity-focused B.E. student in Artificial Intelligence and Machine Learning with an 8.3 CGPI and hands-on experience building Python security CLIs, FastAPI-based scanning platforms, and ML-backed applications. Interested in cybersecurity internships spanning SOC and blue-team operations, Linux, networking, threat detection, and secure software development. Open to internship and fresher opportunities across remote, onsite, and hybrid environments.",
        size=10.0,
        after=2,
        line=1.08,
    )

    add_section_heading(document, "Education")
    p = document.add_paragraph()
    set_spacing(p, after=0.5, line=1.03)
    add_text(
        p,
        "Bachelor of Engineering (B.E.) in Artificial Intelligence and Machine Learning",
        size=10.4,
        bold=True,
    )
    add_body_paragraph(
        document,
        "Bahubali College of Engineering, Shravanabelagola, Karnataka | Expected July 2027 | CGPI: 8.3 / 10",
        size=9.9,
        after=0.4,
        line=1.03,
        color=MUTED,
    )
    add_body_paragraph(
        document,
        "Class XII | 76%  |  Class X | 76%",
        size=9.8,
        after=1.5,
        line=1.03,
        color=MUTED,
    )

    add_section_heading(document, "Experience and Recognition")
    add_bullet(
        document,
        "Received internship and job opportunity recognition from Dashin Technosoft, reflecting early industry interest in cybersecurity-oriented development.",
        size=9.8,
        after=0.5,
        line=1.03,
    )
    add_bullet(
        document,
        "Participated in hackathons and technical project events, strengthening teamwork, rapid problem-solving, and solution building under time constraints.",
        size=9.8,
        after=0.5,
        line=1.03,
    )
    add_bullet(
        document,
        "Active learner in cybersecurity, Linux, networking, and AI/ML with public GitHub projects and certification-backed technical growth.",
        size=9.8,
        after=1.4,
        line=1.03,
    )

    add_section_heading(document, "Technical Skills")
    add_body_paragraph(
        document,
        "Cybersecurity: Network security, SOC and blue-team fundamentals, Linux security, threat analysis, phishing detection, malware basics, supply-chain scanning",
        size=9.8,
        after=0.5,
        line=1.04,
    )
    add_body_paragraph(
        document,
        "Programming and AI/ML: Python, JavaScript, HTML/CSS, React, FastAPI, TensorFlow/Keras, machine learning, AI automation",
        size=9.8,
        after=0.5,
        line=1.04,
    )
    add_body_paragraph(
        document,
        "Tools and Platforms: Git, GitHub, Kali Linux, Linux, Docker, SQLite, Oracle Cloud, VS Code, Google Play Console",
        size=9.8,
        after=1.6,
        line=1.04,
    )

    add_section_heading(document, "Selected Projects")
    add_project(
        document,
        "FastSec / CipherShield Cyber Scanner",
        "Python, FastAPI, JavaScript, HTML/CSS, SQLite, SSL/TLS",
        [
            "Built a full-stack cybersecurity scanner for SSL/TLS certificate inspection and domain intelligence with FastAPI backend services and browser-based frontend flows.",
            "Implemented JWT authentication, scan history, watchlist monitoring, WHOIS and DNS lookups, and alert-oriented workflows for practical security analysis.",
        ],
    )
    add_project(
        document,
        "Safe Security Scanner",
        "Python, CLI, SBOM, pip-audit, supply-chain security",
        [
            "Developed a Python CLI to analyze local repositories, Git URLs, PyPI packages, and SBOMs for risky code patterns, hard-coded secrets, and dependency issues.",
            "Added JSON reporting and Kali-friendly execution paths to support first-pass software supply-chain risk triage in developer and security workflows.",
        ],
    )
    add_project(
        document,
        "Exploit Verifier",
        "Python, Docker, CVE validation, defensive security",
        [
            "Created a defensive exposure-verification CLI that checks public CVE indicators using Docker-isolated benign probes and evidence-based confidence scoring.",
            "Designed the tool for authorized validation without weaponized payload execution, making it useful for blue-team triage, lab validation, and documentation workflows.",
        ],
    )
    add_project(
        document,
        "Blood Group Prediction",
        "Flask, TensorFlow/Keras, OpenCV, HTML",
        [
            "Built an end-to-end ML-backed web application that predicts one of eight labels from uploaded fingerprint images using a trained CNN model.",
            "Connected model inference to a browser workflow with image upload, prediction, and confidence display to demonstrate applied AIML deployment.",
        ],
    )

    add_section_heading(document, "Certifications")
    add_body_paragraph(
        document,
        "Google Foundation of Cybersecurity; Google Play It Safe: Manage Security Risks; Google Tools of the Trade: Linux and SQL; Google Connect and Protect: Networks and Network Security; Cisco Introduction to Cybersecurity",
        size=9.7,
        after=1.4,
        line=1.03,
    )

    add_section_heading(document, "Additional Information")
    add_bullet(
        document,
        "Languages: Kannada, English, Hindi.",
        size=9.8,
        after=0.5,
        line=1.03,
    )
    add_bullet(
        document,
        "Target roles: Cybersecurity Internship, SOC / Blue Team, Software plus Security.",
        size=9.8,
        after=0.5,
        line=1.03,
    )
    add_bullet(
        document,
        "Work preference: Internship or full-time fresher opportunities across remote, onsite, and hybrid environments.",
        size=9.8,
        after=0.5,
        line=1.03,
    )

    document.save(output_path)


if __name__ == "__main__":
    project_root = Path(__file__).resolve().parents[1]
    output = project_root / "public" / "Mokshith-HS-Resume.docx"
    build_resume(output)
    print(f"Wrote {output}")
